from pathlib import Path
import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re
import subprocess
import tempfile
import os
import docx.oxml.shared
from docx.oxml import OxmlElement
import shutil
import uuid

def convert_mermaid_to_image(mermaid_code):
    """将 Mermaid 代码转换为图片"""
    # 创建临时文件
    with tempfile.NamedTemporaryFile(suffix='.mmd', delete=False, mode='w', encoding='utf-8') as f:
        # 添加主题和样式设置
        mermaid_config = """
%%{init: {'theme': 'default', 'themeVariables': { 'fontSize': '16px', 'fontFamily': '宋体' }}}%%
"""
        f.write(mermaid_config + mermaid_code)
        mmd_file = f.name
    
    # 创建输出图片文件
    png_file = mmd_file.replace('.mmd', '.png')
    
    try:
        # 使用 mmdc 命令转换，设置统一的图片大小和背景
        subprocess.run([
            'mmdc',
            '-i', mmd_file,
            '-o', png_file,
            '-w', '800',  # 设置宽度
            '-H', '600',  # 设置高度
            '-b', 'transparent',  # 设置透明背景
            '-s', '3',  # 设置缩放比例
            '-c', 'config.json'  # 使用配置文件
        ], check=True)
        return png_file
    except subprocess.CalledProcessError as e:
        print(f"转换流程图失败: {e}")
        return None
    finally:
        # 清理临时文件
        if os.path.exists(mmd_file):
            os.unlink(mmd_file)

def create_mermaid_config():
    """创建 Mermaid 配置文件"""
    config = {
        "theme": "default",
        "themeVariables": {
            "fontSize": "16px",
            "fontFamily": "宋体",
            "primaryColor": "#1f77b4",
            "primaryTextColor": "#000000",
            "primaryBorderColor": "#1f77b4",
            "lineColor": "#1f77b4",
            "secondaryColor": "#ff7f0e",
            "tertiaryColor": "#2ca02c"
        },
        "flowchart": {
            "curve": "basis",
            "padding": 15,
            "nodeSpacing": 50,
            "rankSpacing": 50
        }
    }
    
    with open('config.json', 'w', encoding='utf-8') as f:
        import json
        json.dump(config, f, indent=2)

def process_mermaid(doc, mermaid_code):
    """处理 Mermaid 流程图"""
    # 转换 Mermaid 代码为图片
    png_file = convert_mermaid_to_image(mermaid_code)
    if png_file and os.path.exists(png_file):
        try:
            # 添加图片到文档
            doc.add_picture(png_file, width=Inches(6))  # 先插入图片
            
            # 设置图片居中
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加图片说明（可选）
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run("图 X-X 流程图")
            caption_run.font.name = '宋体'
            caption_run.font.size = Pt(10.5)
        finally:
            # 清理临时图片文件
            os.unlink(png_file)

def set_document_styles(doc):
    """设置文档样式"""
    # 设置默认字体
    styles = doc.styles
    style = styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 设置标题样式
    for i in range(1, 5):
        style = styles[f'Heading {i}']
        style.font.name = '黑体'
        style.font.size = Pt(16 - i)  # 标题字号递减
        if i == 1:
            style.font.bold = True
    
    # 设置列表样式
    style = styles['List Bullet']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    style = styles['List Number']
    style.font.name = '宋体'
    style.font.size = Pt(12)

def set_document_format(doc, project_name):
    """设置文档格式"""
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
        
        # 添加页眉
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = f"{project_name}投标文件"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加页脚
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "第 "
        # 当前页码
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        footer_para.add_run(" 页，共 ")
        # 总页数
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.text = 'NUMPAGES'
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        footer_para.add_run(" 页")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def process_table(md_table, doc):
    """处理 Markdown 表格"""
    lines = md_table.strip().split('\n')
    if len(lines) < 3:  # 至少需要表头、分隔行和一行数据
        return
    
    # 计算列数
    header_cells = lines[0].strip('|').split('|')
    col_count = len(header_cells)
    
    # 创建表格
    table = doc.add_table(rows=1, cols=col_count)
    table.style = 'Table Grid'
    
    # 添加表头
    header_row = table.rows[0]
    for i, cell in enumerate(header_cells):
        header_row.cells[i].text = cell.strip()
        # 设置表头格式
        for paragraph in header_row.cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '黑体'
    
    # 添加数据行
    for line in lines[2:]:  # 跳过表头和分隔行
        cells = line.strip('|').split('|')
        if len(cells) == col_count:
            row = table.add_row()
            for i, cell in enumerate(cells):
                row.cells[i].text = cell.strip()
                # 设置单元格格式
                for paragraph in row.cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = '宋体'

def convert_md_to_word(md_file):
    """将Markdown文件转换为Word文档"""
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档格式
    project_name = Path(md_file).parent.name
    set_document_format(doc, project_name)
    
    # 处理Markdown内容
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 处理表格
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            process_table('\n'.join(table_lines), doc)
            continue
        
        # 处理标题
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            # 移除标题中的加粗标记
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line.lstrip('#').strip())
            if level == 1:
                # 一级标题作为文档标题
                doc.add_heading(text, level=0)
            else:
                # 其他级别的标题
                doc.add_heading(text, level=level-1)
        
        # 处理列表
        elif line.startswith(('- ', '* ', '+ ')):
            # 移除列表标记
            text = line[2:].strip()
            # 移除加粗标记
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
        
        # 处理数字列表
        elif re.match(r'^\d+\.', line):
            # 移除数字和点
            text = re.sub(r'^\d+\.', '', line).strip()
            # 移除加粗标记
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(style='List Number')
            p.add_run(text)
        
        # 处理普通段落
        elif line:
            # 移除加粗标记
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            p = doc.add_paragraph()
            p.add_run(text)
        
        i += 1
    
    # 保存文档：目标文件名与 md 同名（.docx），先写入临时文件再替换，遇到被占用时退化为带唯一后缀的文件
    parent = Path(md_file).parent
    parent.mkdir(parents=True, exist_ok=True)
    output_file = Path(md_file).with_suffix('.docx')

    temp_path = None
    try:
        tf = tempfile.NamedTemporaryFile(dir=str(parent), suffix='.docx', delete=False)
        temp_path = Path(tf.name)
        tf.close()

        # 保存到临时文件
        doc.save(str(temp_path))

        # 尝试原子替换目标文件
        try:
            os.replace(str(temp_path), str(output_file))
            saved_path = output_file
        except PermissionError:
            # 目标被占用（常见于 Windows），改为生成带唯一后缀的备份文件
            alt_name = parent / f"{output_file.stem}_{uuid.uuid4().hex}.docx"
            shutil.move(str(temp_path), str(alt_name))
            saved_path = alt_name
            print(f"目标文件被占用，已生成备用文件：{saved_path}")

        print(f"已生成 Word 文档：{saved_path}")
        return Path(saved_path)
    finally:
        # 清理残留临时文件（如果存在）
        try:
            if temp_path and temp_path.exists():
                temp_path.unlink()
        except Exception:
            pass

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        print(convert_md_to_word(sys.argv[1]))
    else:
        print("请传入md文件路径，例如：python md_to_word.py data/output/项目名/项目名_完整投标文件.md")
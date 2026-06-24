from pathlib import Path
import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re
import subprocess
import tempfile
import os
import docx.oxml.shared
from docx.oxml import OxmlElement
import shutil
import uuid
from export.flowchart_renderer import build_flowchart_svg, parse_flowchart_payload, parse_mermaid_flowchart

IMAGE_MARKDOWN_RE = re.compile(r'^!\[([^\]]*)\]\(([^)\s]+)(?:\s+"[^"]*")?\)\s*$')
HTML_BREAK_RE = re.compile(r'<br\s*/?>', re.IGNORECASE)
TABLE_DIVIDER_RE = re.compile(r'^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$')


def _to_chapter_number(value):
    if not value:
        return None
    value = str(value).strip()
    if value.isdigit():
        return value
    digits = {
        "零": 0,
        "一": 1,
        "二": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
    }
    if value == "十":
        return "10"
    if value.startswith("十") and len(value) == 2:
        return str(10 + digits.get(value[1], 0))
    if value.endswith("十") and len(value) == 2:
        return str(digits.get(value[0], 0) * 10)
    if "十" in value and len(value) == 3:
        return str(digits.get(value[0], 0) * 10 + digits.get(value[2], 0))
    if value in digits:
        return str(digits[value])
    return None


def extract_chapter_number(text):
    match = re.search(r'第\s*([一二三四五六七八九十百\d]+)\s*章', text or "")
    if match:
        return _to_chapter_number(match.group(1))
    match = re.match(r'^(\d+)[、．.]\s*', text or "")
    if match:
        return match.group(1)
    return None


def convert_mermaid_to_image(mermaid_code):
    """将 Mermaid 代码转换为图片，优先用 mmdc，失败则用纯文本框替代"""
    try:
        with tempfile.NamedTemporaryFile(suffix='.mmd', delete=False, mode='w', encoding='utf-8') as f:
            f.write(mermaid_code)
            mmd_file = f.name
        png_file = mmd_file.replace('.mmd', '.png')
        subprocess.run(
            ['mmdc', '-i', mmd_file, '-o', png_file, '-w', '800', '-b', 'white'],
            check=True, timeout=15, capture_output=True,
        )
        if os.path.exists(mmd_file):
            os.unlink(mmd_file)
        return png_file
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        try:
            if os.path.exists(mmd_file):
                os.unlink(mmd_file)
        except Exception:
            pass
        return None


def render_mermaid_as_text_box(doc, mermaid_code):
    """当 mmdc 不可用时，将 Mermaid 代码渲染为 Word 中的文本框图表"""
    lines = [line.strip() for line in mermaid_code.strip().split('\n') if line.strip()]
    nodes = []
    for line in lines:
        cleaned = re.sub(r'[\[\]{}()|>]', '', line)
        cleaned = re.sub(r'--.*?>', ' → ', cleaned)
        cleaned = re.sub(r'---', ' → ', cleaned)
        cleaned = re.sub(r'--', ' → ', cleaned)
        cleaned = cleaned.strip()
        if cleaned and not cleaned.startswith(('flowchart', 'graph', 'sequenceDiagram', 'gantt', 'classDiagram', '%%')):
            nodes.append(cleaned)

    if not nodes:
        return

    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("【流程图】\n")
    run.font.name = '黑体'
    run.font.size = Pt(10)
    run.bold = True

    for node in nodes[:15]:
        run = p.add_run(f"\n{node}")
        run.font.name = '宋体'
        run.font.size = Pt(9)


def render_flowchart_as_table(doc, chart):
    nodes = chart.get("nodes") or []
    if not nodes:
        return
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = p.add_run(f"【{chart.get('title') or '流程图'}】\n")
    title_run.font.name = '黑体'
    title_run.font.size = Pt(10.5)
    title_run.bold = True
    flow_text = "  →  ".join(node.get("text") or "" for node in nodes)
    run = p.add_run(flow_text)
    run.font.name = '宋体'
    run.font.size = Pt(10)


def _insert_svg_chart(doc, chart):
    svg = build_flowchart_svg(chart)
    try:
        import cairosvg
    except Exception:
        return False
    png_file = None
    try:
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
            png_file = f.name
        cairosvg.svg2png(bytestring=svg.encode("utf-8"), write_to=png_file)
        doc.add_picture(png_file, width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except Exception:
        return False
    finally:
        try:
            if png_file and os.path.exists(png_file):
                os.unlink(png_file)
        except Exception:
            pass


def add_flowchart_caption(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title or "流程图")
    run.font.name = '宋体'
    run.font.size = Pt(10.5)


def process_structured_flowchart(doc, payload):
    try:
        chart = parse_flowchart_payload(payload)
    except Exception:
        return False
    if not _insert_svg_chart(doc, chart):
        render_flowchart_as_table(doc, chart)
    add_flowchart_caption(doc, chart.get("title") or "流程图")
    return True


def create_mermaid_config():
    """创建 Mermaid 配置文件"""
    config = {
        "theme": "default",
        "themeVariables": {
            "fontSize": "16px",
            "fontFamily": "宋体",
            "primaryColor": "#1f77b4",
            "primaryTextColor": "#000000",
            "primaryBorderColor": "#1f77b4",
            "lineColor": "#1f77b4",
            "secondaryColor": "#ff7f0e",
            "tertiaryColor": "#2ca02c"
        },
        "flowchart": {
            "curve": "basis",
            "padding": 15,
            "nodeSpacing": 50,
            "rankSpacing": 50
        }
    }
    
    with open('config.json', 'w', encoding='utf-8') as f:
        import json
        json.dump(config, f, indent=2)

def process_mermaid(doc, mermaid_code):
    """处理 Mermaid 流程图：先转结构化流程图，再兜底 mmdc/text box，避免代码入 Word"""
    try:
        chart = parse_mermaid_flowchart(mermaid_code)
        if chart.get("nodes"):
            if not _insert_svg_chart(doc, chart):
                render_flowchart_as_table(doc, chart)
            add_flowchart_caption(doc, chart.get("title") or "流程图")
            return
    except Exception:
        pass

    png_file = convert_mermaid_to_image(mermaid_code)
    if png_file and os.path.exists(png_file):
        try:
            doc.add_picture(png_file, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        finally:
            try:
                os.unlink(png_file)
            except Exception:
                pass
        return
    render_mermaid_as_text_box(doc, mermaid_code)

def set_document_styles(doc):
    """设置文档样式"""
    styles = doc.styles
    style = styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)

    for i in range(1, 5):
        style = styles[f'Heading {i}']
        style.font.name = '黑体'
        style.font.size = Pt(16 - i)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True

    try:
        style = styles['List Bullet']
        style.font.name = '宋体'
        style.font.size = Pt(12)
    except KeyError:
        pass

    try:
        style = styles['List Number']
        style.font.name = '宋体'
        style.font.size = Pt(12)
    except KeyError:
        pass

def set_document_format(doc, project_name):
    """设置文档格式"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = f"{project_name}投标文件"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.font.size = Pt(9)
            run.font.name = '宋体'

        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run("第 ")
        run.font.size = Pt(9)
        run2 = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run2._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run2._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run2._r.append(fldChar2)
        run3 = footer_para.add_run(" 页 共 ")
        run3.font.size = Pt(9)
        run4 = footer_para.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run4._r.append(fldChar3)
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = 'NUMPAGES'
        run4._r.append(instrText2)
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run4._r.append(fldChar4)
        run5 = footer_para.add_run(" 页")
        run5.font.size = Pt(9)


def generate_cover_page(doc, cover_data, facts):
    """根据招标文件封面要求生成投标文件封面"""
    if not cover_data or not cover_data.get("cover_lines"):
        return generate_default_cover(doc, facts)

    for _ in range(3):
        doc.add_paragraph()

    for line_info in cover_data["cover_lines"]:
        text = line_info.get("text", "")
        style = line_info.get("style", "normal")
        placeholder = line_info.get("placeholder", "none")

        if placeholder and placeholder != "none":
            value = facts.get(placeholder, "")
            if value:
                text = value

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)

        if style == "title":
            run.font.size = Pt(26)
            run.font.name = '方正小标宋简体'
            run.bold = True
            p.space_after = Pt(20)
        elif style == "subtitle":
            run.font.size = Pt(18)
            run.font.name = '黑体'
            run.bold = True
            p.space_after = Pt(12)
        else:
            run.font.size = Pt(14)
            run.font.name = '宋体'
            p.space_after = Pt(8)

    for _ in range(2):
        doc.add_paragraph()

    if cover_data.get("cover_notes"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cover_data["cover_notes"])
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def generate_default_cover(doc, facts):
    """生成默认封面"""
    for _ in range(4):
        doc.add_paragraph()

    project_name = facts.get("project_name", "投标文件")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project_name)
    run.font.size = Pt(26)
    run.font.name = '方正小标宋简体'
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("投 标 文 件")
    run.font.size = Pt(36)
    run.font.name = '方正小标宋简体'
    run.bold = True

    for _ in range(3):
        doc.add_paragraph()

    info_lines = [
        ("投标人：", facts.get("bidder_name", "________________")),
        ("法定代表人或授权代表：", "________________"),
        ("日期：", facts.get("bid_date", "____年____月____日")),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}{value}")
        run.font.size = Pt(14)
        run.font.name = '宋体'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("（此封面需加盖投标人公章）")
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

def clean_markdown_text(text):
    """清除文本中的 Markdown 语法标记"""
    if not text:
        return ""
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'__(.*?)__', r'\1', text)
    text = re.sub(r'_(.*?)_', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    text = re.sub(r'~~(.*?)~~', r'\1', text)
    return text.strip()


def is_table_divider(line):
    return bool(TABLE_DIVIDER_RE.match(line or ""))


def parse_table_row(line):
    if "|" not in (line or ""):
        return None
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return cells if len(cells) >= 2 else None


def is_markdown_table_start(lines, index):
    if index + 1 >= len(lines):
        return False
    return parse_table_row(lines[index]) is not None and is_table_divider(lines[index + 1])


def normalize_table_cell_text(text):
    text = HTML_BREAK_RE.sub("\n", text or "")
    text = text.replace("&nbsp;", " ")
    return clean_markdown_text(text)


def normalize_table_cells(cells, col_count):
    if len(cells) == col_count:
        return cells
    if len(cells) < col_count:
        return cells + [""] * (col_count - len(cells))
    return cells[: col_count - 1] + [" | ".join(cells[col_count - 1 :])]


def process_table(md_table, doc):
    """处理 Markdown 表格"""
    lines = [line for line in md_table.strip().split('\n') if line.strip()]
    if len(lines) < 3:
        return

    header_row = parse_table_row(lines[0])
    if not header_row or not is_table_divider(lines[1]):
        return

    header_cells = [normalize_table_cell_text(c) for c in header_row]
    col_count = len(header_cells)

    table = doc.add_table(rows=1, cols=col_count)
    table.style = 'Table Grid'

    header_row = table.rows[0]
    for i, cell in enumerate(header_cells):
        header_row.cells[i].text = cell
        for paragraph in header_row.cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '黑体'
                run.font.size = Pt(10.5)

    for line in lines[2:]:
        if not line.strip() or is_table_divider(line):
            continue
        parsed_cells = parse_table_row(line)
        if not parsed_cells:
            continue
        cells = [normalize_table_cell_text(c) for c in normalize_table_cells(parsed_cells, col_count)]
        if len(cells) == col_count:
            row = table.add_row()
            for i, cell in enumerate(cells):
                row.cells[i].text = cell
                for paragraph in row.cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(10.5)


def add_image_placeholder(doc, line, chapter_no=None, image_index=None):
    match = IMAGE_MARKDOWN_RE.match((line or "").strip())
    if not match:
        return False

    caption = clean_markdown_text(match.group(1).strip()) or "章节配图"
    target = match.group(2).strip()
    if chapter_no and image_index:
        figure_no = f"{chapter_no}-{image_index}"
    else:
        figure_no = str(getattr(doc, "_ai_bidding_image_count", 0) + 1)
    setattr(doc, "_ai_bidding_image_count", getattr(doc, "_ai_bidding_image_count", 0) + 1)

    inserted_real_image = False
    if not target.startswith("image-plan://"):
        image_path = Path(target)
        if image_path.exists() and hasattr(doc, "add_picture"):
            try:
                doc.add_picture(str(image_path), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                inserted_real_image = True
            except Exception:
                inserted_real_image = False

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    status = "图片待补充" if target.startswith("image-plan://") or not inserted_real_image else "图片已插入"
    run = p.add_run(f"图 {figure_no} {caption}（{status}）")
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    return True


def add_template_paragraph(doc, line):
    text = clean_markdown_text(line.strip())
    if not text:
        doc.add_paragraph()
        return
    heading_match = re.match(r'^#+\s*(.+)$', text)
    numbered_heading_match = None
    if len(text) <= 24 and not re.search(r'[。；;：:]', text):
        numbered_heading_match = re.match(r'^(?:[一二三四五六七八九十]+|\d+)[、．.]\s*(.{2,20})$', text)
    if heading_match or numbered_heading_match:
        text = clean_markdown_text(heading_match.group(1) if heading_match else numbered_heading_match.group(1))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(16)
        run.bold = True
        return

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)


def convert_md_to_word(md_file, cover_data=None, facts=None):
    """将Markdown文件转换为Word文档"""
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()

    filename_stem = Path(md_file).stem
    project_name = filename_stem.replace('_bid_document', '').replace('_完整投标文件', '')
    if facts and facts.get("project_name"):
        project_name = facts["project_name"]

    set_document_format(doc, project_name)

    generate_cover_page(doc, cover_data, facts or {"project_name": project_name})
    
    # 处理Markdown内容
    lines = md_content.split('\n')
    i = 0
    in_locked_template = False
    current_chapter_no = None
    current_chapter_image_index = 0
    while i < len(lines):
        line = lines[i].strip()

        if line == '<!-- locked_template -->':
            in_locked_template = True
            i += 1
            continue
        if line == '<!-- /locked_template -->':
            in_locked_template = False
            i += 1
            continue
        if in_locked_template:
            add_template_paragraph(doc, lines[i])
            i += 1
            continue

        if IMAGE_MARKDOWN_RE.match(line):
            current_chapter_image_index += 1
            add_image_placeholder(doc, line, current_chapter_no or "1", current_chapter_image_index)
            i += 1
            continue

        # 处理结构化流程图代码块
        if line.startswith('```flowchart-json') or line.startswith('```flowchart'):
            i += 1
            flowchart_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                flowchart_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            process_structured_flowchart(doc, '\n'.join(flowchart_lines))
            continue

        # 处理 Mermaid 代码块
        if line.startswith('```mermaid'):
            i += 1
            mermaid_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                mermaid_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            process_mermaid(doc, '\n'.join(mermaid_lines))
            continue

        # 跳过其他代码块
        if line.startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            if i < len(lines):
                i += 1
            continue

        # 处理表格，兼容有/无首尾竖线的 Markdown 表格
        if is_markdown_table_start(lines, i):
            table_lines = []
            while i < len(lines) and (parse_table_row(lines[i]) or is_table_divider(lines[i])):
                table_lines.append(lines[i])
                i += 1
            process_table('\n'.join(table_lines), doc)
            continue
        
        # 处理标题
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = clean_markdown_text(line.lstrip('#').strip())
            chapter_no = extract_chapter_number(text)
            if chapter_no:
                current_chapter_no = chapter_no
                current_chapter_image_index = 0
            if level == 1:
                doc.add_heading(text, level=0)
            else:
                doc.add_heading(text, level=min(level - 1, 4))

        # 处理列表
        elif line.startswith(('- ', '* ', '+ ')):
            text = clean_markdown_text(line[2:].strip())
            try:
                p = doc.add_paragraph(style='List Bullet')
            except KeyError:
                p = doc.add_paragraph()
                text = '• ' + text
            run = p.add_run(text)
            run.font.name = '宋体'
            run.font.size = Pt(12)

        # 处理数字列表
        elif re.match(r'^\d+[\.\、]', line):
            text = clean_markdown_text(re.sub(r'^\d+[\.\、]\s*', '', line).strip())
            try:
                p = doc.add_paragraph(style='List Number')
            except KeyError:
                p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = '宋体'
            run.font.size = Pt(12)

        # 处理普通段落
        elif line:
            text = clean_markdown_text(line)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(text)
            run.font.name = '宋体'
            run.font.size = Pt(12)
        
        i += 1
    
    # 保存文档：目标文件名与 md 同名（.docx），先写入临时文件再替换，遇到被占用时退化为带唯一后缀的文件
    parent = Path(md_file).parent
    parent.mkdir(parents=True, exist_ok=True)
    output_file = Path(md_file).with_suffix('.docx')

    temp_path = None
    try:
        tf = tempfile.NamedTemporaryFile(dir=str(parent), suffix='.docx', delete=False)
        temp_path = Path(tf.name)
        tf.close()

        # 保存到临时文件
        doc.save(str(temp_path))

        # 尝试原子替换目标文件
        try:
            os.replace(str(temp_path), str(output_file))
            saved_path = output_file
        except PermissionError:
            # 目标被占用（常见于 Windows），改为生成带唯一后缀的备份文件
            alt_name = parent / f"{output_file.stem}_{uuid.uuid4().hex}.docx"
            shutil.move(str(temp_path), str(alt_name))
            saved_path = alt_name

        return Path(saved_path)
    finally:
        # 清理残留临时文件（如果存在）
        try:
            if temp_path and temp_path.exists():
                temp_path.unlink()
        except Exception:
            pass

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        print(convert_md_to_word(sys.argv[1]))
    else:
        print("请传入md文件路径，例如：python md_to_word.py data/output/项目名/项目名_完整投标文件.md")
